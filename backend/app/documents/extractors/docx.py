"""DOCX extractor (python-docx).

Note: the DOCX format has no intrinsic page boundaries, so `page_count` is
reported as 1 when text is present (a documented limitation).
"""

from __future__ import annotations

import io

from app.documents.exceptions import ExtractionError
from app.documents.extractors.base import ExtractionResult, TextExtractor
from app.documents.extractors.registry import register_extractor
from app.documents.model import DocumentFormat


@register_extractor(DocumentFormat.DOCX)
class DocxExtractor(TextExtractor):
    """Extracts paragraph and table text from a .docx file."""

    @property
    def document_format(self) -> DocumentFormat:
        return DocumentFormat.DOCX

    def extract(self, content: bytes) -> ExtractionResult:
        try:
            import docx
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ExtractionError("python-docx is not installed") from exc

        try:
            document = docx.Document(io.BytesIO(content))
        except Exception as exc:
            raise ExtractionError(f"Could not open DOCX: {exc}") from exc

        lines = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                lines.append("\t".join(cells))

        text = "\n".join(lines)
        return ExtractionResult(
            text=text,
            page_count=1 if text.strip() else 0,
            page_texts=[text] if text.strip() else [],
        )
